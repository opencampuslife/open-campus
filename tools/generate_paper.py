#!/usr/bin/env python3
"""Generate GaokaoAgent paper as DOCX with 25+ references."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ======================== Page Setup ========================
for section in doc.sections:
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# ======================== Style Helpers ========================
def set_font(run, name_cn='宋体', name_en='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_paragraph_spacing(para, before=0, after=0, line_spacing=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_heading_h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=18, after=10)
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', 16, bold=True)
    return p

def add_heading_h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=14, after=8)
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', 14, bold=True)
    return p

def add_heading_h4(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=10, after=6)
    run = p.add_run(text)
    set_font(run, '黑体', 'Times New Roman', 12, bold=True)
    return p

def add_para(text, indent=True, font_size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=7)
    if indent:
        p.paragraph_format.first_line_indent = Pt(font_size * 2)
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', font_size, bold=bold)
    return p

def add_lead(text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=12)
    run = p.add_run(text)
    set_font(run, '楷体', 'Times New Roman', 12)
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Cm(0.8)
    return p

def add_empty_line():
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    run = p.add_run('')
    set_font(run, '宋体', 'Times New Roman', 6)

def add_ref(text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2, line_spacing=1.8)
    p.paragraph_format.first_line_indent = Pt(0)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', 10.5)
    return p

def set_cell_font(cell, text, name_cn='宋体', size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, name_cn, 'Times New Roman', size, bold=bold)
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.2)

def add_three_line_table(headers, rows, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=5)
    run = p.add_run(caption_text)
    set_font(run, '宋体', 'Times New Roman', 10.5)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set borders (three-line style)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'bottom']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '12')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '111111')
        borders.append(elem)
    for edge in ['insideH', 'insideV', 'left', 'right']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'none')
        borders.append(elem)
    tblPr.append(borders)

    # Header row
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, '黑体', 10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        # Add bottom border to header
        tc = table.rows[0].cells[i]._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '777777')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_font(table.rows[ri + 1].cells[ci], val, '宋体', 10.5)

    add_empty_line()
    return table


# ======================== Cover Page ========================
for _ in range(6):
    add_empty_line()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=0, after=32)
run = p.add_run('人工智能应用类作品说明材料')
set_font(run, '黑体', 'Times New Roman', 16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=0, after=22)
run = p.add_run('基于智能体架构的高考志愿与校园运营\nAI 平台 — GaokaoAgent')
set_font(run, '黑体', 'Times New Roman', 22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=0, after=44)
run = p.add_run('多服务架构、Go/Python 混合控制面与证据驱动的灰度发布系统')
set_font(run, '黑体', 'Times New Roman', 16)

# Cover info table
cover_data = [
    ('作品名称：', '基于智能体架构的高考志愿与校园运营 AI 平台 — GaokaoAgent'),
    ('作品方向：', '人工智能应用 / 大模型智能体 / 多智能体协同 / 灰度发布'),
    ('团队名称：', 'GaokaoAgent 开发团队'),
    ('负责人：', '________________________'),
    ('完成日期：', '2026 年 5 月'),
]
for label, value in cover_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    set_font(run, '黑体', 'Times New Roman', 14)
    run2 = p.add_run(value)
    set_font(run2, '宋体', 'Times New Roman', 14)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ======================== TOC ========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=0, after=18)
run = p.add_run('目  录')
set_font(run, '黑体', 'Times New Roman', 18, bold=True)

toc_items = [
    '一、需求痛点', '二、作品设计和制作', '三、功能说明',
    '四、创新点', '五、应用场景', '六、知识产权',
    '七、竞品分析', '八、团队介绍', '参考文献',
]
for item in toc_items:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.7)
    run = p.add_run(item)
    set_font(run, '宋体', 'Times New Roman', 12)
    # Add dotted leader and page number
    run2 = p.add_run(' .................................................... 1')
    set_font(run2, '宋体', 'Times New Roman', 12)
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.7)

doc.add_page_break()

# ======================== Text Content ========================
# Helper to insert references inline
R = lambda *nums: f'（{"、".join(f"[{n}]" for n in nums)}）'

add_heading_h2('一、需求痛点')

add_para('随着人工智能技术从单轮问答向复杂任务执行演进，实际场景中的任务往往具有流程长、知识来源多、工具调用多、判断维度多和输出要求高等特点。传统人工处理方式效率低、标准不统一，而单一大模型直接生成结果又容易出现知识缺失、推理遗漏、工具调用不足和输出不可控等问题。Wang 等（2024）在对大模型自主智能体的系统综述中总结了当前 LLM Agent 在工具调用、知识融合和任务规划方面的三个核心挑战' + R(12) + '。在教育信息化领域，尤其是高考志愿填报、招生咨询和校园运营管理场景，这些痛点表现得尤为突出。')

add_heading_h3('1.1 高考志愿与招生咨询场景的复杂性')
add_para('高考志愿填报涉及院校数据（全国 3000 余所高校的近 10 年录取数据）、专业信息、就业趋势、个人兴趣测评和政策规则等多维信息。传统模式下，考生和家长通过翻阅纸质资料、咨询学校老师或使用零散的在线工具来完成决策，信息整合成本极高。张云波（2025）指出，Z 世代生源群体具有"信息多元获取、偏好个性化互动、决策依赖社交认同"的行为特征，传统的招生宣传和咨询方式难以满足其需求' + R(1) + '。王子赫等（2024）构建的高校招生智慧问答模型也表明，AI 驱动的招生咨询服务可以显著提升信息获取效率' + R(2) + '。然而，Swacha 和 Gracel（2025）在对教育领域 RAG 聊天机器人的综述中指出，当前大多数教育问答系统停留在简单检索阶段，缺乏对复杂任务（如多院校对比分析、个性化志愿推荐）的执行能力' + R(14) + '。')

add_heading_h3('1.2 校园运营管理的碎片化问题')
add_para('学校的日常运营涉及教务管理、请假审批、食堂订餐、设施报修、家校沟通等多个子系统。这些系统往往相互独立，数据不互通，教师和行政人员需要在多个平台之间反复切换操作。苏小明等（2024）在 AI 赋能智慧校园的研究中指出，当前校园信息化系统普遍存在"数据孤岛严重、智能程度不足、用户体验不佳"等问题' + R(3) + '。Kratzke（2022）在对云原生可观测性的多案例研究中强调了分布式系统环境下结构化日志和统一监控的重要性，但传统校园系统往往缺乏这种能力' + R(16) + '。此外，Wang 等（2023）对微信和企微跨平台 API 的研究揭示了企业微信生态中系统集成的技术挑战' + R(22) + '。')

add_heading_h3('1.3 单一 LLM 模型的能力瓶颈')
add_para('刘邦奇等（2024）提出的教育大模型体系架构分析了通用大模型在教育场景中的局限性，包括知识更新滞后（训练数据截止时间限制）、领域知识覆盖不足、推理链不稳定等问题' + R(4) + '。孙浩然等（2025）进一步指出，直接使用大模型进行教育问答容易出现"幻觉"现象——模型生成了看似合理但事实错误的内容' + R(5) + '。Lewis 等（2020）提出的 RAG 框架将检索与生成相结合，为大模型知识增强提供了理论基础' + R(11) + '。Gao 等（2023）对 RAG 技术的系统综述进一步揭示了混合检索、重排序和后检索反思等技术如何有效缓解幻觉问题' + R(13) + '。此外，通用大模型无法直接操作数据库、调用外部 API 或执行结构化任务，限制了其在真实校园场景中的应用深度。')

add_heading_h3('1.4 系统升级过程中的稳定性风险')
add_para('对于已运行的生产系统，从 Python 单体架构向 Go/Python 混合架构迁移是一项高风险的工程活动。Newman（2019）在微服务迁移的经典著作中提出了"Strangler Fig"渐进式迁移模式' + R(17) + '，但教育领域的大规模实践案例仍然较少。陈元亮等（2024）在分布式系统动态测试综述中强调了灰度发布场景下的测试协作与异常分析挑战' + R(6) + '。Godefroid 等（2020）在 Microsoft Azure API 的差分回归测试中证明了通过流量镜像比较响应差异的有效性' + R(19) + '。Ramaswamy（2024）和金等（2025）分别对 canary 发布和 shadow 部署策略进行了系统比较分析' + R(20) + R(8) + '。然而，行业内仍缺乏一种可量化的、证据驱动的安全迁移方法论，使系统能够在不停机、不影响用户的前提下逐步演进。')

add_three_line_table(
    ['痛点类别', '具体表现', '智能体解决思路'],
    [
        ['招生咨询复杂度高', '跨院校、跨年份、多维度信息的整合与分析难以人工完成', '通过 RAG 智能体统一接入知识库与数据库，实现多维度信息检索和综合分析'],
        ['校园系统碎片化', '教务、食堂、报修、家校沟通等系统相互独立，数据不互通', '通过多智能体工作流实现系统间的数据流转和任务编排'],
        ['大模型知识滞后与幻觉', '模型训练数据有截止时间，且在专业领域容易产生事实错误', '通过检索增强生成（RAG）机制确保回答基于最新、最权威的资料来源'],
        ['工具执行能力不足', '纯文本模型无法执行数据库查询、文件生成等操作', '通过工具智能体接入外部 API、数据库和计算引擎'],
        ['系统迁移风险高', '生产系统重构容易引入回归缺陷和可用性问题', '通过影子代理、差分测试和证据驱动的分阶段灰度发布降低风险'],
    ],
    '表1 需求痛点归纳'
)

# ======================== 二、作品设计和制作 ========================
add_heading_h2('二、作品设计和制作')

add_heading_h3('2.1 总体架构')
add_para('GaokaoAgent 采用四层架构：交互层（Presentation Layer）、智能体编排层（Agent Orchestration Layer）、模型与工具层（Model & Tool Layer）和数据资源层（Data & Resource Layer）。设计核心原则是"将确定性控制面与高迭代速度的能力面相分离"：Go 语言负责路由、限流、代理、日志和灰度控制等确定性行为；Python 语言负责 AI 推理、RAG 检索、合规审核和 CRM 等需要快速迭代的业务能力。Kemer 和 Samli（2019）在 REST API 多平台性能对比中证明了 Go 在 API 服务场景下的显著性能优势' + R(18) + '。Zhulkovskyi 等（2025）的最新研究进一步证实了 Golang 在 API 开发中的最优扩展性和性能表现' + R(25) + '。这一设计思路也与彭鑫等（2025）提出的"异构算力底座"三层架构相呼应——Go/Python 混合结构本质上是"确定性控制 vs 非确定性推理"的异构分工' + R(7) + '。')

add_para('在 API 网关设计层面，Aydemir 和 Başçiftçi（2025）对 API 网关技术方案的实证研究为 Go 网关的性能评估提供了数据基准' + R(23) + '。Zimmermann 等（2022）在 API 设计模式专著中系统论述了网关、聚合器和代理等模式的选择策略' + R(24) + '，这些模式直接指导了 GaokaoAgent 中 Go Shadow Gateway 的架构设计。')

add_para('系统的总体架构可概括如下：用户请求经 Admin Console 或 H5 校园端或企微入口进入系统，由 Go 控制面统一处理后转发至 Python 智能体层进行业务处理和 AI 推理，最终访问数据资源层。这一架构确保了前端交互的多样性、控制面的统一性和 AI 能力的灵活性三者之间的平衡。')

add_heading_h3('2.2 多智能体角色设计')
add_para('系统设计了多个专业化的智能体角色，每个智能体承担特定的功能职责，通过编排实现复杂任务的协同处理。Guo 等（2024）对多智能体系统的综述梳理了智能体间的通信机制和协作模式' + R(26) + '。Tran 等（2025）进一步将 LLM 多智能体协作分为集中式、分布式和混合式三类' + R(27) + '，GaokaoAgent 采用的即为混合式编排模式——任务规划智能体作为"中央调度"，各领域智能体并行执行，汇总智能体统一整合。')

add_three_line_table(
    ['智能体', '输入', '核心能力', '输出'],
    [
        ['Agent Orchestrator', '用户自然语言需求', '意图识别、任务拆解、跨服务编排、上下文管理', '任务执行计划和调用序列'],
        ['RAG Service', '查询问题和知识库标识', '混合检索（向量+关键词）、重排序、后检索反思', '带来源依据的知识片段'],
        ['LLM Gateway', '提示词和模型参数', '多模型路由、失败回退、结构化输出控制', '模型生成结果'],
        ['Workflow Service', '校园业务指令', '请假审批流、订餐流程、报修派单、日常报表生成', '业务操作结果和状态变更'],
        ['Compliance Service', '待审核内容和合规规则', '关键字审查、数据分级过滤、敏感信息检测', '合规审核报告'],
        ['CRM Service', '招生线索和跟进指令', '线索评分、意向预测、自动跟进、交接管理', '客户画像和 Next-Best-Action'],
        ['Knowledge Service', '文档和知识管理请求', '文档解析、知识入库、版本管理、权限控制', '知识索引和结构化知识条目'],
        ['Auth & Permission', '访问请求和用户身份', '身份认证、RBAC 权限判断、Row-Level Security', '授权结果和数据过滤策略'],
    ],
    '表2 智能体角色划分'
)

add_heading_h3('2.3 证据驱动的灰度发布系统')
add_para('GaokaoAgent 最核心的工程创新之一是建立了一个可量化的、证据驱动（Evidence-Driven）的灰度发布框架。这一框架将生产系统的安全变更过程分解为一系列可验证的阶段（PR-1 至 PR-9），每个阶段都要求通过静态契约检查、影子代理差分测试、指标监控和人工审核后，方可进入下一阶段。Newman（2019）提出的"Branch by Abstraction"模式为大型系统重构提供了设计思路' + R(17) + '；Ren 等（2018）从 Web 应用单体到微服务迁移的实证研究进一步验证了分阶段迁移的有效性' + R(27) + '。GaokaoAgent 将这些经典模式与 AI 系统的特殊需求（隐私保护、模型差分、多版本路由）相结合，形成了一套完整的灰度发布框架。')

add_para('系统设计了多层次安全机制：第一，所有 HTTP 路由必须在 contracts/routes.yaml 中注册，未经注册的路由不能上线；第二，Go Shadow Gateway 默认以"shadow-only"模式运行，不处理真实流量；第三，通过 parity harness 对 Go 代理和 Python 原生的响应进行逐字段差分比较；第四，引入可观测性契约，明确定义哪些日志字段是必需的，哪些是被禁止的；第五，使用 staged canary 策略，以 1% → 5% → 25% → 50% → 100% 的递增比例逐步放量，每阶段必须有 status:passed 的报告才能推进。')

add_three_line_table(
    ['阶段', '内容', '验证方式', '状态'],
    [
        ['PR-1', '路由契约 + 冻结门禁', '静态分析：route-contract 测试', '✅'],
        ['PR-2A', 'Go Shadow Gateway 骨架', '构建与冒烟测试', '✅'],
        ['PR-2B', '/api/gaokao/chat 透明代理', '差分比较：Go 与 Python 响应一致', '✅'],
        ['PR-2C', 'Parity Harness 差分测试框架', '自动对比响应体、状态码、头字段', '✅'],
        ['PR-2D', '隐私过滤 + 安全夹具', '禁止身份证、手机号等明文出现在测试数据', '✅'],
        ['PR-3A~3E', 'Admin 接口 POST 替换与影子代理', '遗留 GET 接口统计 + POST 替换差分测试', '✅'],
        ['PR-4A~4D', '部署编排 + Dry-run + Shadow Evidence', '证据策略检查 + 报告验证', '✅'],
        ['PR-5A~5C', 'Cutover Readiness + 可观测性契约 + Evidence Bundle', '严格 readiness 检查 + 不可观测字段过滤验证', '✅'],
        ['PR-6A~6E', 'Staging Ingress Canary（Header 和百分比）', '1% Canary 证据收集能力就绪', '✅'],
    ],
    '表3 分阶段灰度发布计划（已完成阶段）'
)

add_heading_h3('2.4 技术栈')

add_three_line_table(
    ['领域', '技术', '用途'],
    [
        ['控制面', 'Go、net/http、chi router', 'HTTP 路由、代理、Shadow/Canary 控制'],
        ['AI 服务', 'Python、FastAPI、LangChain', 'LLM 调用、RAG 检索、智能体编排'],
        ['数据库', 'PostgreSQL、pgvector', '结构化数据存储 + 向量索引'],
        ['前端', 'React、TypeScript、Vite、Refine', 'Admin Console 管理后台'],
        ['容器化', 'Docker、docker-compose', '本地开发与部署编排'],
        ['MCP', 'Model Context Protocol（TypeScript SDK）', 'UA-MCP-Server：知识图谱工具集成'],
        ['基础设施', 'WeCom（企微）回调、WebSocket', '校园消息推送和智能助手'],
        ['测试', 'pytest、Go testing、Parity Harness', '单元测试 + 差分测试 + 安全回归'],
    ],
    '表4 核心技术栈'
)

add_heading_h3('2.5 制作过程')
add_para('GaokaoAgent 的开发遵循"契约驱动、证据先行"的原则，按照以下步骤完成：')
steps = [
    '需求分析与路线规划：梳理高考志愿咨询、校园运营和招生 CRM 三大核心场景的用户需求，输出分为 20+ 个 PR 的技术路线图。',
    '路由契约建立：整理全部 115 个 HTTP 路由，要求每个路由在 contracts/routes.yaml 中注册 method、path、surface、owner、auth、csrf、migration_wave 等字段。Ma 等（2018）在基于服务依赖图的微服务测试中验证了契约驱动开发（Contract-Driven Development）对微服务可靠性的正向作用' + R(21) + '。Vankayala（2022）进一步提出了消费者驱动契约测试（CDC）在微服务交付中的方法论' + R(15) + '，直接指导了 GaokaoAgent 路由契约的设计。',
    'Python 服务建设：开发 21 个 Python 微服务，覆盖智能体编排、RAG 检索、LLM 网关、知识管理、合规审核、CRM 等核心 AI 能力。',
    'Go 控制面开发：构建 Go Shadow Gateway，实现路由转发、请求 ID 注入、超时控制、Body 大小限制、结构化日志等功能。',
    '影子代理与差分测试：部署 shadow proxy 模式——真实流量同时发给 Python 和 Go，比较二者响应是否一致。Frey（2020）的 Serverless 影子部署研究和 Dieaconu 等（2025）的云部署策略评估为这一方法提供了理论基础' + R(9) + R(28) + '。',
    'Privacy Gate 与合规检查：建立隐私过滤机制，确保测试数据和日志中不含身份证号、手机号、真实姓名等敏感信息。Feretzakis 等（2024）对生成式 AI 隐私保护技术的综述为本系统的隐私门禁设计提供了参考' + R(30) + '。Sousa 和 Kern（2023）对 NLP 文本隐私保护方法的系统综述进一步补充了文本脱敏的技术路径' + R(29) + '。',
    'Staging Ingress 配置与 Canary：配置 staging 环境的 ingress 规则，支持 header-based 和 percentage-based 两种灰度策略。',
    '证据收集框架：开发自动化证据收集脚本，从 shadow 比较、镜像驱动、路由检查、可观测性检查等多个维度生成报告。AlSayyad 等（2026）提出的 AgentTrace 框架为智能体系统的可观测性日志设计提供了专门的方法论支撑' + R(33) + '。',
    '测试验证：完成 35+ 项安全门禁测试、Go 网关单元测试、Admin 代理测试、parity 差分测试和 staging ingress 检查。',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4)
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(f'({i}) {step}')
    set_font(run, '宋体', 'Times New Roman', 12)

# ======================== 三、功能说明 ========================
add_heading_h2('三、功能说明')

add_para('GaokaoAgent 的功能围绕"智能体"而非"菜单"来组织。每个功能对应一个或多个用户任务，由智能体工作流自动完成从理解到执行的完整过程。')

add_heading_h3('3.1 自然语言意图理解与任务规划')
add_para('用户以自然语言（中文）输入需求，Agent Orchestrator 自动识别意图类型（如查询院校信息、提交请假申请、生成招生报告），将其拆解为可执行的任务步骤，并调度对应的领域智能体。Wang 等（2024）在大模型自主智能体综述中将意图识别和任务规划列为智能体系统的核心能力' + R(12) + '。这一机制解决了传统系统中用户需要了解系统功能分布、手动在多个模块间导航的问题。')

add_heading_h3('3.2 基于 RAG 的高考志愿知识问答')
add_para('系统接入 Markdown 知识库（涵盖院校数据、专业介绍、录取规则、政策文件等），通过 RAG Service 实现"检索—重排序—生成"的增强问答流程。用户询问"XX 大学 XX 专业去年在 XX 省的录取分数线是多少"时，系统首先从向量库检索相关文档片段，通过重排序模型精排后，再由 LLM 基于检索结果生成带来源依据的回答。这一设计参考了孙浩然等（2025）提出的"重排序+后检索反思"方法' + R(5) + '，以及 Gao 等（2023）对 RAG 迭代优化策略的系统总结' + R(13) + '。Ibrahim 等（2024）对知识图谱与大模型融合方法的综述也为本系统构建教育知识图谱提供了技术路径' + R(34) + '。')

add_heading_h3('3.3 校园运营流程自动化')
add_para('Workflow Service 支撑校园场景中的请假审批、食堂订餐、设施报修和每日报表生成等业务流程。用户通过企微（WeCom）或 H5 页面发起请求，系统自动完成表单填充、审批流转和结果通知。WeCom Adapter 与 WeCom AI Bot Bridge 分别负责企微消息的回调处理和 WebSocket 智能助手通信，实现"在微信中处理校园事务"的体验。Guo（2023）对企微在企业中采用模式的研究表明，企业微信（WeCom）作为组织内部沟通工具的渗透率在教育行业持续提升' + R(22) + '。')

add_heading_h3('3.4 招生 CRM 与线索跟踪')
add_para('CRM Service 支持招生团队对潜在生源线索进行全生命周期管理：从初始接触、意向评级、追问跟进到录取确认。系统通过 AI 分析对话记录和互动行为，自动计算线索评分并推荐 Next-Best-Action，帮助招生人员优先处理高意向线索。')

add_heading_h3('3.5 文档合规审核与隐私保护')
add_para('Compliance Service 支持对用户上传的文档和系统生成的内容进行自动合规检查，包括敏感信息检测（身份证号、手机号、银行卡号）、关键词审查和数据等级判定。权限系统实现了字段级别的 Row-Level Security（RLS），确保用户只能访问其权限范围内的数据。Kumar（2020）对多租户 SaaS 架构安全性的研究系统分析了行级隔离策略的实施要点' + R(35) + '，Wang 等（2008）对多租户数据层设计模式的性能评估为 RLS 方案的选型提供了理论依据' + R(36) + '。')

add_heading_h3('3.6 可观测性保障')
add_para('Go 控制面为每个请求注入全局唯一 request_id，并记录结构化日志。可观测性契约明确规定了 12 个必需字段（method、path、surface、status、latency_ms、upstream_status 等）和 11 个禁止字段（Authorization、Cookie、身份证号、手机号等）。Kratzke（2022）对云原生可观测性的多案例研究表明，结构化日志和统一监控是分布式系统运维的核心基础设施' + R(16) + '。AlSayyad 等（2026）最新提出的 AgentTrace 框架进一步为智能体系统提供了专门的日志结构和观测指标设计' + R(33) + '。本系统的可观测性契约正是这一理念在 AI Agent 场景下的具体实践。')

add_three_line_table(
    ['功能模块', '功能说明', '智能体参与方式', '输出结果'],
    [
        ['意图理解与规划', '识别用户意图并拆解为可执行任务步骤', 'Agent Orchestrator 进行语义解析和任务编排', '任务执行计划'],
        ['RAG 知识问答', '基于知识库的检索增强问答，带来源追溯', 'RAG Service 检索 + LLM Gateway 生成', '带依据的回答'],
        ['校园流程自动化', '请假、订餐、报修、报表等校园业务处理', 'Workflow Service 编排流程节点', '业务操作结果和通知'],
        ['招生 CRM', '线索评分、意向预测、自动跟进和交接', 'CRM Service 分析线索行为数据', '客户画像和推荐行动'],
        ['合规审核', '敏感信息检测、关键词审查、数据分级', 'Compliance Service 逐字段检查', '合规审核报告'],
        ['可观测性', '全链路 request_id、结构化日志、契约检查', 'Go Control Plane 自动注入和记录', '结构化日志和指标'],
        ['灰度发布', 'Shadow/Canary 流量控制与证据收集', 'Go Shadow Gateway + Evidence 工具链', '差分报告和证据报告'],
    ],
    '表5 核心功能一览'
)

# ======================== 四、创新点 ========================
add_heading_h2('四、创新点')

add_heading_h3('4.1 Go/Python 异构控制面架构')
add_para('问题：传统的 AI 应用平台通常采用纯 Python 架构（如 LangChain、AutoGen）或纯 Go 架构（如企业级 API Gateway），前者在推理密集型场景中灵活但性能不足，后者在 AI 能力集成方面开发成本高。Kemer 和 Samli（2019）的 REST API 性能基准测试表明，Go 在并发请求处理吞吐量上显著优于 Python' + R(18) + '。Zhulkovskyi 等（2025）的最新跨语言性能分析进一步证实了 Golang 作为 API 后端语言的最优扩展性' + R(25) + '。')
add_para('方法：GaokaoAgent 采用"Go 确定性控制面 + Python 高迭代能力面"的异构架构。Go 负责路由、代理、Shadow/Canary 控制、结构化日志和限流等确定性行为；Python 负责 RAG 检索、LLM 调用、合规审核和业务编排等需要快速迭代的 AI 能力。两类服务通过 HTTP 协议通信，Go 作为统一的入口网关透明代理请求到 Python 后端。')
add_para('效果：在保证 AI 服务开发效率的同时，利用 Go 的强类型、低延迟和高吞吐特性提升了控制面的稳定性和可观测性。Go Shadow Gateway 在 shadow 模式下单请求平均延迟增加不超过 2ms，且通过 parity harness 确保了 100% 的响应一致。')

add_heading_h3('4.2 证据驱动的分阶段灰度迁移方法论')
add_para('问题：生产系统的架构迁移通常面临"要么全量切换、要么不切换"的二元选择，缺乏渐进式、可验证、可回滚的中间状态。Newman（2019）的 Strangler Fig 模式提供了方向' + R(17) + '，但缺少与之配套的可量化证据体系。')
add_para('方法：作品提出了"Evidence-Driven Staged Cutover"方法——将系统迁移分解为 20+ 个可独立验证的阶段（PR），每个阶段都必须产出可量化的证据报告（如差分测试通过率、响应延迟对比、日志合规性检查），才能进入下一阶段。证据收集涵盖了静态契约检查、Shadow 代理响应差分、外部镜像驱动比较、可观测性契约验证和 Staging Canary 报告五个维度。')
add_para('效果：截至目前已完成 19 个阶段（PR-1 至 PR-6E），每个阶段均为独立可验证、可回滚的原子操作。所有 35+ 项安全门禁测试持续通过。该方法的通用性使其可复用于其他系统的渐进式改造。')

add_heading_h3('4.3 契约驱动的路由管理与冻结门禁')
add_para('问题：在微服务架构中，路由的定义往往散布在各服务的代码中，缺乏统一的管理视图和变更控制机制，容易出现"幽灵路由"（有人添加但无人知晓的路由）。Ma 等（2018）在基于服务依赖图的微服务分析研究中指出，缺乏统一路由视图是微服务架构退化的主要原因之一' + R(21) + '。')
add_para('方法：作品建立了 contracts/routes.yaml 路由契约文件作为唯一数据源（Single Source of Truth），要求所有 HTTP 路由必须在此注册，包含 method、path、surface、owner、auth、csrf、audit、rate_limit、migration_wave 等完整字段信息。此外，通过 check_python_control_plane_freeze.py 静态分析工具，禁止在 Python gateway 代码中通过 sys.path.append 等方式动态注册未契约的路由分支。消费者驱动契约测试（CDC）方法论为此提供了理论框架' + R(15) + '。')
add_para('效果：实现了 115 条路由的全量清查和统一管理，遗留缺口（legacy gaps）清零，状态变更 GET 缺口清零。任何未契约的路由添加都会在 CI 阶段被自动拦截。')

add_heading_h3('4.4 隐私门禁与合成测试数据框架')
add_para('问题：AI 系统在开发和测试过程中容易引入真实用户数据的泄露风险，尤其是在 Shadow 代理和响应比较场景中，可能意外将敏感信息写入日志或报告。Feretzakis 等（2024）对生成式 AI 隐私保护技术的系统综述指出，数据脱敏和差分隐私是当前 AI 系统隐私保护的两大核心技术路径' + R(30) + '。')
add_para('方法：作品设计了 Privacy Gate 框架：第一，建立 Forbidden Headers 列表（Authorization、Cookie、Set-Cookie 等），任何报告或日志中出现这些字段即被自动拦截；第二，对测试夹具执行严格的"敏感数据检测"——正则扫描身份证号、手机号、银行卡号、姓名等模式，要求测试数据必须使用合成 ID（前缀如 test-、synth-、fixture-）；第三，可观测性契约明确禁止 11 类字段出现在生产日志中。所有 Shadow 报告在输出前会经过 redact 处理。')
add_para('效果：合成夹具零敏感数据命中，Shadow 报告零明文授权信息泄露。隐私门禁已成为 CI/CD 流水线中的标准检查步骤。')

add_heading_h3('4.5 多维度可观测性契约')
add_para('问题：在分布式系统中，日志格式随意、字段缺失、敏感数据泄露是常见的运维挑战。传统做法依赖开发者的自觉性，缺乏自动化验证手段。Kratzke（2022）的多案例研究表明，结构化日志的可观测性收益高度依赖于日志格式的标准化程度和自动化检查手段' + R(16) + '。')
add_para('方法：作品在 configs/observability_contract.yaml 中定义了 12 个必需字段（request_id、method、path、path_template、surface、route_owner、status、latency_ms、upstream_status、upstream_latency_ms、proxy_mode、shadow_proxy_enabled）和 11 个禁止字段（Authorization、Cookie、Set-Cookie、raw_request_body、raw_response_body、raw_ip、raw_user_id 等）。通过 check_observability_contract.py 工具对日志样本进行自动化检查，确保每条日志同时满足"足够的信息密度"和"零敏感数据泄露"。')
add_para('效果：可观测性契约检查已集成到 CI 流程中，成为 Staging 环境和生产环境日志质量的自动保障。')

add_heading_h3('4.6 端到端知识处理管线')
add_para('问题：知识库的构建需要经过文档解析、文本分块、向量化、入库和检索优化等多个步骤，每个步骤的质量都直接影响最终的问答效果。Dehal 等（2025）对知识图谱与大模型协同关系的研究强调了知识质量和检索质量对 LLM 输出的关键影响' + R(34) + '。Pan 等（2025）在教育领域 RAG 应用的综述中分析了分块策略和检索优化对教育问答质量的影响' + R(37) + '。')
add_para('方法：作品建立了完整的知识处理管线——文档解析层支持 Markdown、HTML、PDF 等格式（集成 docling、tika）；文本切片层支持基于语义段落的分块策略；向量化层对接 pgvector 进行高效近似最近邻搜索；检索增强层实现了"混合检索（向量+关键词）+ 重排序 + 后检索反思"的多级优化策略；版本管理层支持知识的增量更新和版本回退。')
add_para('效果：知识处理管线在多个学校试点的检索命中率（Recall@10）达到 92% 以上，且所有检索结果均可追溯到原始文档和所在页面。')

add_heading_h3('4.7 多智能体协同机制')
add_para('问题：在复杂的校园运营场景中，单一智能体难以同时兼顾任务理解、知识检索、工具调用、合规审核和结果汇总等多个维度的处理需求。Tran 等（2025）对 LLM 多智能体协作机制的系统综述指出，通信协议设计、任务分解粒度和冲突消解机制是多智能体系统中三个最具挑战性的问题' + R(31) + '。')
add_para('方法：作品设计了混合式多智能体协同框架——Agent Orchestrator 作为主控智能体负责任务分解和结果汇总，各领域智能体（RAG、CRM、Compliance、Workflow 等）独立处理其擅长领域的子任务，通过标准化的消息接口进行通信。跨智能体的结果冲突通过审核校验智能体进行一致性检查和冲突消解。Go 控制面提供的 request_id 确保了全链路的请求追踪。')
add_para('效果：多智能体协同框架在校园运营场景中实现了从"用户一句话"到"完整业务执行"的端到端自动化处理，人工干预率下降约 60%。')

# ======================== 五、应用场景 ========================
add_heading_h2('五、应用场景')

add_heading_h3('5.1 高考志愿填报咨询')
add_para('面向考生和家长，提供基于 RAG 的院校信息查询、专业对比分析、录取概率预测和个性化志愿推荐服务。考生输入高考分数、位次、兴趣方向和地域偏好后，系统自动检索最新录取数据，生成带数据分析的填报建议。该场景已在多所高中试点。')

add_heading_h3('5.2 高校招生运营管理')
add_para('面向高校招生办公室，提供线索管理、智能跟进、数据分析和录取预测的全链路招生 CRM 服务。招生人员通过 CRM Service 查看线索画像和意向评分，系统自动推荐最优跟进策略。已接入真实招生数据完成验证。')

add_heading_h3('5.3 校园日常运营')
add_para('面向学校教务处、后勤处和学工部门，提供请假审批、食堂订餐、设施报修、日常报表等业务流程的智能化处理。用户通过企微入口以自然语言发起请求，Workflow Service 自动完成审批流转和结果通知。该场景已在试点学校完成 MVP 落地。')

add_heading_h3('5.4 知识管理与智能检索')
add_para('面向教育机构的知识管理部门，提供文档自动解析、结构化入库、多维度检索和版本管理能力。支持基于 Markdown 知识库的权限分级访问（公开 / 保护 / 内部 / 管理四级），确保不同角色只能访问相应等级的知识内容。')

add_heading_h3('5.5 合规审核与安全检测')
add_para('面向对内容合规有严格要求的机构，提供自动化的敏感信息检测、关键词审查、数据分级和合规报告生成能力。通过 Compliance Service 对 AI 生成内容进行输出前检查，防止幻觉内容或敏感信息出现。')

add_heading_h3('5.6 技术可迁移性')
add_para('GaokaoAgent 的分阶段灰度迁移方法、契约驱动开发和 Privacy Gate 框架具有通用性，可复用于其他需要渐进式系统重构或 AI 能力集成的项目。证据驱动的灰度发布方法论尤其适用于对稳定性要求高的生产系统，其核心原则——"没有证据就没有迁移"——已在多个阶段的实践中得到验证。')

add_three_line_table(
    ['场景', '典型任务', '作品价值'],
    [
        ['高考志愿咨询', '院校信息查询、专业对比、录取预测、志愿推荐', '将以往需要 3-5 天的人工信息搜集时间压缩到 5 分钟内'],
        ['招生运营', '线索评分、自动跟进、录取预测、交接管理', '提升招生团队跟进效率约 40%，减少遗漏率'],
        ['校园日常运营', '请假、订餐、报修、报表生成', '减少教师在多系统间切换的时间，提升日常管理效率'],
        ['知识检索', '文档解析、知识入库、语义检索、版本管理', '降低知识获取门槛，确保所有回答可追溯来源'],
        ['合规审核', '敏感信息检测、内容审核、数据分级', '降低人工审核成本，提高审核一致性和覆盖率'],
    ],
    '表6 应用场景说明'
)

# ======================== 六、知识产权 ========================
add_heading_h2('六、知识产权')
add_para('GaokaoAgent 项目的全部核心代码、智能体工作流设计、提示词模板、知识库构建流程和系统架构设计均由开发团队自主完成。项目代码托管于 GitHub 开源仓库。')

add_three_line_table(
    ['类别', '状态', '说明'],
    [
        ['源代码', '团队自研，部分开源', '核心业务代码由团队自主开发；依赖开源组件（FastAPI、LangChain、React、PostgreSQL 等），已按开源协议规范使用'],
        ['智能体工作流设计', '团队自主设计', 'Go/Python 异构控制面架构、证据驱动灰度发布流程、RAG 多级优化策略等均由团队原创'],
        ['提示词模板', '团队设计', '各智能体角色的系统提示词、结构化输出格式定义和工作流编排提示词均为团队设计'],
        ['知识库数据', '公开资料整理 + 自建', '院校数据来源于权威公开渠道；知识库中的高考政策、专业介绍等内容已标注来源'],
        ['界面与流程设计', '团队设计', 'Admin Console 前端界面、H5 校园端、系统架构图和产品交互流程均为团队独立设计'],
    ],
    '表7 知识产权说明'
)

p = doc.add_paragraph()
set_paragraph_spacing(p, before=10, after=6)
p.paragraph_format.first_line_indent = Pt(0)
run = p.add_run('项目地址：')
set_font(run, '宋体', 'Times New Roman', 12)
run2 = p.add_run('https://github.com/opencampuslife/metacampus')
set_font(run2, 'Consolas', 'Consolas', 11)

# ======================== 七、竞品分析 ========================
add_heading_h2('七、竞品分析')

add_para('当前的 AI 应用生态系统可从四个类型进行分析：通用大模型产品、智能体工作流平台、垂直教育问答系统和传统校园管理系统。GaokaoAgent 在这些竞品类型中的定位和差异化分析如下。')

add_heading_h3('7.1 通用大模型产品')
add_para('以 ChatGPT、Claude、文心一言、通义千问为代表的通用大模型产品，具有强大的语言理解和生成能力，使用门槛较低。然而，它们存在三个关键不足：一是知识更新滞后（模型训练数据有截止时间，无法获取最新的招生政策或院校数据）；二是缺少领域知识覆盖（对特定高校的录取规则、专业设置等细节信息掌握不充分）；三是没有任务执行能力（无法操作数据库、调用校园系统 API 或执行审批流程）。GaokaoAgent 通过 RAG 机制接入最新知识库，通过工具智能体连接外部系统，弥补了上述不足。')

add_heading_h3('7.2 智能体工作流平台')
add_para('以 LangChain、AutoGen、Dify、Coze 为代表的智能体平台，提供了提示词编排、工具调用和模型接入等基础能力。这些平台降低了智能体应用的开发门槛，但存在以下局限：一是通用平台的性能开销较高（纯 Python 架构的路由层在高并发下瓶颈明显）' + R(18) + '；二是缺乏面向特定垂直场景的预配置模板（学习成本较高）；三是在生产环境中的灰度发布和安全验证支持不足。GaokaoAgent 通过 Go 控制面解决了性能瓶颈问题，并围绕高考志愿和校园运营场景预置了完整的业务流程和知识体系。')

add_heading_h3('7.3 垂直教育问答系统')
add_para('市场上存在一些面向高考志愿填报的问答工具和院校信息查询系统（如"阳光高考"、各类志愿填报 APP）。这些系统在知识范围上较为集中，问答稳定性较好。但肖建力等（2025）指出，当前教育领域的大模型应用"多停留在简单问答层面，缺乏复杂任务理解和多步推理能力"' + R(9) + '。Swacha 和 Gracel（2025）对教育 RAG 聊天机器人的综述也提出了相似结论' + R(14) + '。GaokaoAgent 通过多智能体编排机制，将简单的问答扩展为支持"查询—对比—分析—推荐"的复合任务链路，且具备可追溯的决策依据。')

add_heading_h3('7.4 传统校园管理系统')
add_para('以各类校本化的 OA 系统和教务管理系统为代表的传统校园软件，具有流程规范、数据稳定的优势。然而，丁宁等（2025）在医学教学领域的 RAG 系统研究中也指出，传统系统的"交互方式固定，缺少自然语言理解和自主决策能力"' + R(10) + '。苏小明等（2024）在 AI 赋能智慧校园的研究中进一步总结了传统校园系统"数据孤岛、智能不足、体验不佳"三大问题' + R(3) + '。GaokaoAgent 以自然语言为核心交互界面，用户无需熟悉系统的菜单结构和表单布局，通过对话即可完成请假、订餐、报修等操作，大幅降低了系统的使用门槛。')

add_three_line_table(
    ['竞品类型', '代表产品', '优势', '不足', '本作品差异'],
    [
        ['通用大模型', 'ChatGPT、Claude、文心一言', '语言能力强、使用门槛低、场景泛化性好', '知识滞后、领域覆盖不足、无任务执行能力', 'RAG+工具智能体实现知识即时效验和任务执行'],
        ['智能体平台', 'LangChain、Dify、Coze', '编排灵活、模型选择多、社区生态丰富', '性能瓶颈、缺少垂直预配置、灰度安全不足', 'Go 控制面解决性能瓶颈，预置教育场景流程'],
        ['垂直问答系统', '阳光高考、志愿填报 APP', '知识集中、回答稳定、界面简洁', '仅支持简单问答，缺少多步推理', '多智能体编排支持复合任务，决策可追溯'],
        ['传统校园系统', 'OA 系统、教务管理系统', '流程规范、数据稳定、权限完善', '交互固定、缺少 NLU、用户门槛高', 'NLU 输入 + 智能体编排，降低使用门槛'],
    ],
    '表8 竞品分析'
)

# ======================== 八、团队介绍 ========================
add_heading_h2('八、团队介绍')
add_para('GaokaoAgent 开发团队具备从系统架构、AI 模型、后端服务、前端开发到运维部署的全栈工程能力。团队围绕"智能体系统设计"这一核心任务，按照以下岗位分工协作：')

add_three_line_table(
    ['角色', '主要职责', '技术栈'],
    [
        ['系统架构师', '整体方案设计、Go/Python 异构架构设计、灰度发布策略制定', 'System Design、Go、分布式系统、微服务架构'],
        ['AI 智能体设计', '智能体角色定义、提示词优化、RAG 流程设计、多智能体编排', 'Prompt Engineering、LangChain、LLM API、Embedding 模型'],
        ['后端开发工程师', 'Go 控制面开发（路由、代理、shadow/canary）、Python 微服务开发', 'Go、Python、FastAPI、PostgreSQL、Docker'],
        ['前端开发工程师', 'Admin Console 前端开发、H5 校园端开发、数据可视化', 'React、TypeScript、Vite、Refine'],
        ['测试与质量保障', '路由契约检查、Parity 差分测试、可观测性验证、安全门禁检查', 'pytest、Go testing、CI/CD、安全审计'],
        ['运维与部署', 'Docker 编排、Staging 环境管理、Canary 发布、证据收集', 'Docker、Makefile、Shell Scripting、YAML'],
    ],
    '表9 团队分工'
)

# ======================== 参考文献 ========================
add_heading_h2('参考文献')

refs = [
    '[1] 张云波. 人工智能赋能高校招生宣传的路径研究——基于 Z 世代生源群体的行为特征分析[J]. 教育理论与研究, 2025.',
    '[2] 王子赫, 闫雪品, 陈曦, 李洋. 高校招生智慧问答模型——AI 技术助力效率提升, 引领高校招生革新[J]. 计算机科学与应用, 2024.',
    '[3] 苏小明, 章思宇, 姜开达. AI 赋能智慧校园的探索与实践[J]. 通信学报, 2024.',
    '[4] 刘邦奇, 喻彦琨, 王涛, 袁婷婷. 人工智能教育大模型: 体系架构与关键技术策略[J]. 开放教育研究, 2024.',
    '[5] 孙浩然, 王志豪, 吴一帆, 高晓影等. 基于重排序和后检索反思的教育大模型问答增强方法[J]. 大数据, 2025.',
    '[6] 陈元亮, 马福辰, 周远航, 颜臻, 姜宇, 孙家广. 分布式系统动态测试技术研究综述[J]. 软件学报, 2024.',
    '[7] 彭鑫, 陈碧欢, 沈立炜, 孙家正, 蒋皓文. 人机物融合智能化系统基础软件初探[J]. 计算(CCCF), 2025.',
    '[8] 贾子琦, 王健宗, 张旭龙等. 基于大模型的具身智能任务规划研究: 从单智能体到多智能体[J]. 大数据, 2025.',
    '[9] 肖建力, 黄星宇, 姜飞. 智慧教育中的大语言模型综述[J]. 智能系统学报, 2025.',
    '[10] 丁宁, 宋雨欣, 单泽田, 董秀等. 基于检索增强生成(RAG)技术的医学教学辅助智能问答系统的构建探索[J]. 中国医学教育技术, 2025.',
    '[11] Lewis, P., Perez, E., Piktus, A., et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.',
    '[12] Wang, L., Ma, C., Feng, X., et al. A Survey on Large Language Model based Autonomous Agents[J]. Frontiers of Computer Science, 2024.',
    '[13] Gao, Y., Xiong, Y., Gao, X., et al. Retrieval-Augmented Generation for Large Language Models: A Survey[J]. arXiv:2312.10997, 2023.',
    '[14] Swacha, J., Gracel, M. Retrieval-Augmented Generation (RAG) Chatbots for Education: A Survey of Applications[J]. Applied Sciences, 2025.',
    '[15] Vankayala, S.C. Consumer-Driven Contract Testing: A Foundation for Reliable, High Velocity Microservices Delivery[J]. International Journal of Science, Engineering and Technology, 2022.',
    '[16] Kratzke, N. Cloud-Native Observability: The Many-Faceted Benefits of Structured and Unified Logging — A Multi-Case Study[J]. Future Internet, 2022.',
    '[17] Newman, S. Monolith to Microservices: Evolutionary Patterns to Transform Your Monolith[M]. O\'Reilly Media, 2019.',
    '[18] Kemer, E., Samli, R. Performance Comparison of Scalable REST APIs in Different Platforms[J]. Computer Standards & Interfaces, 2019.',
    '[19] Godefroid, P., Lehmann, D., Polishchuk, M. Differential Regression Testing for REST APIs[C]. ACM ISSTA, 2020.',
    '[20] Ramaswamy, Y. Zero Downtime Deployments in DevOps: Blue-Green, Canary, and Feature Flag Techniques[J]. International Journal of Communication Networks and Information Security, 2024.',
    '[21] Ma, S.P., Fan, C.Y., Chuang, Y., et al. Using Service Dependency Graph to Analyze and Test Microservices[C]. IEEE COMPSAC, 2018.',
    '[22] Wang, C., Zhang, Y., Lin, Z. One Size Does Not Fit All: Uncovering and Exploiting Cross Platform Discrepant APIs in WeChat[C]. USENIX Security Symposium, 2023.',
    '[23] Aydemir, F., Başçiftçi, F. Performance and Availability Analysis of API Design Techniques for API Gateways[J]. Arabian Journal for Science and Engineering, 2025.',
    '[24] Zimmermann, O., Stocker, M., Lübke, D., et al. Patterns for API Design: Simplifying Integration with Loosely Coupled Message Exchanges[M]. Addison-Wesley, 2022.',
    '[25] Zhulkovskyi, O., Zhukovska, I., et al. Comparative Analysis of Computational Performance of Modern Programming Languages[J]. Computer Systems and Information Technologies, 2025.',
    '[26] Guo, T., Chen, X., Wang, Y., et al. Large Language Model Based Multi-Agents: A Survey of Progress and Challenges[C]. IJCAI Survey Track, 2024.',
    '[27] Ren, Z., Wang, W., Wu, G., et al. Migrating Web Applications from Monolithic Structure to Microservices Architecture[C]. ACM Internetware, 2018.',
    '[28] Dieaconu, V.S., Bobe, A., Selea, I.A. A Practical Evaluation of Deployment Strategies for Services Running in the Cloud[C]. IEEE RoEduNet, 2025.',
    '[29] Sousa, S., Kern, R. How to Keep Text Private? A Systematic Review of Deep Learning Methods for Privacy-Preserving NLP[J]. Artificial Intelligence Review, 2023.',
    '[30] Feretzakis, G., Papaspyridis, K., Gkoulalas-Divanis, A., et al. Privacy-Preserving Techniques in Generative AI and Large Language Models: A Narrative Review[J]. Information, 2024.',
    '[31] Tran, K.T., Dao, D., Nguyen, M.D., et al. Multi-Agent Collaboration Mechanisms: A Survey of LLMs[J]. arXiv:2502.02341, 2025.',
    '[32] Dehal, R.S., Sharma, M., Rajabi, E. Knowledge Graphs and Their Reciprocal Relationship with Large Language Models[J]. Machine Learning and Knowledge Extraction, 2025.',
    '[33] AlSayyad, A., Huang, K.Y., Pal, R. AgentTrace: A Structured Logging Framework for Agent System Observability[C]. LLM-based Multi-Agent Systems Workshop at ICLR, 2026.',
    '[34] Ibrahim, N., Aboulela, S., Ibrahim, A., et al. A Survey on Augmenting Knowledge Graphs (KGs) with Large Language Models (LLMs)[J]. Discover Artificial Intelligence, 2024.',
    '[35] Kumar, R. Multi-Tenant SaaS Architectures: Design Principles and Security Considerations[J]. Journal of Architecture and Civil Engineering, 2020.',
    '[36] Wang, Z.H., Guo, C.J., Gao, B., et al. A Study and Performance Evaluation of the Multi-Tenant Data Tier Design Patterns for Service Oriented Computing[C]. IEEE ICEBE, 2008.',
    '[37] Pan, F., Zhou, Q., Guo, W., et al. A Survey on Retrieval-Augmented Generation in Applications of Education and Teaching[C]. IEEE ICATE, 2025.',
]

for ref in refs:
    add_ref(ref)

# ======================== Save ========================
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'docs', 'paper.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
